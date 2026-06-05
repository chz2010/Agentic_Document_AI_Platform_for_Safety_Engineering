"""Upload parsing and chunking utilities."""

from __future__ import annotations

import csv
from pathlib import Path

import pdfplumber
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from backend.settings import settings


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".markdown", ".csv", ".docx"}


def extract_documents(path: Path) -> list[Document]:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported document type: {suffix}")
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".csv":
        return _extract_csv(path)
    if suffix == ".docx":
        return _extract_docx(path)
    return [Document(page_content=path.read_text(encoding="utf-8", errors="ignore"), metadata={"page": None})]


def chunk_documents(docs: list[Document]) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=["\n## ", "\n# ", "\n\n", "\n", ". ", " ", ""],
    )
    return splitter.split_documents(docs)


def _extract_pdf(path: Path) -> list[Document]:
    docs: list[Document] = []
    with pdfplumber.open(path) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                docs.append(Document(page_content=text, metadata={"page": page_number}))
    return docs


def _extract_csv(path: Path) -> list[Document]:
    rows: list[str] = []
    with path.open("r", encoding="utf-8", errors="ignore", newline="") as handle:
        reader = csv.DictReader(handle)
        for idx, row in enumerate(reader, start=1):
            values = [f"{key}: {value}" for key, value in row.items()]
            rows.append(f"Row {idx}: " + "; ".join(values))
    return [Document(page_content="\n".join(rows), metadata={"page": None})]


def _extract_docx(path: Path) -> list[Document]:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:  # pragma: no cover
        raise ValueError("DOCX parsing requires python-docx") from exc
    doc = DocxDocument(path)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
    return [Document(page_content=text, metadata={"page": None})]
